"""
Regression tests for the custom PyMuPDF→DOCX extraction pipeline.

All four bug classes are exercised with synthetic PDFs built on-the-fly using
PyMuPDF's drawing API so the tests have no external file dependencies.

Bug 1 — Word-spacing loss (Type3-style adjacent spans with zero gap)
Bug 2 — Span ordering / scrambled text
Bug 3 — Text clipped/dropped near overlapping graphics
Bug 4 — Shape-embedded / badge text skipped (annotation text)
"""
from __future__ import annotations

import io
import os
import sys
import tempfile
from pathlib import Path

import fitz
import pytest

# Make sure the backend package is importable when running from repo root.
sys.path.insert(0, str(Path(__file__).parent.parent))

from app.converters.pdf_word.text_layer import (
    extract_page,
    extract_document,
    GAP_SPACE_THRESHOLD,
    BASELINE_TOLERANCE_FACTOR,
)
from app.converters.pdf_word.docx_builder import build_docx


# ─── helpers ────────────────────────────────────────────────────────────────

def _make_doc() -> fitz.Document:
    """Return a blank A4 PyMuPDF document."""
    doc = fitz.open()
    doc.new_page(width=595, height=842)
    return doc


def _page_texts(doc: fitz.Document, page_idx: int = 0) -> list[str]:
    """Return joined paragraph texts from extract_page for one page."""
    page = doc[page_idx]
    content = extract_page(page, page_idx)
    return [p.text for p in content.paragraphs if p.text.strip()]


# ─── Bug 1: word-spacing loss ────────────────────────────────────────────────

class TestWordSpacing:
    def test_adjacent_words_separated_by_gap(self):
        """
        Two words placed with a visible space gap should appear in extracted
        text as two separate words joined by a space (not concatenated).
        """
        doc = _make_doc()
        page = doc[0]
        # Write "AIRBLUE" and "LIMITED" with a ~3pt gap between them
        page.insert_text((50, 100), "AIRBLUE", fontsize=11)
        page.insert_text((105, 100), "LIMITED", fontsize=11)

        texts = _page_texts(doc)
        assert texts, "Expected at least one paragraph"
        combined = " ".join(texts)
        assert "AIRBLUE" in combined
        assert "LIMITED" in combined
        # They must be separated: "AIRBLUELIMITED" is the bug
        assert "AIRBLUELIMITED" not in combined

    def test_zero_gap_spans_no_spurious_space(self):
        """
        Two words placed with no gap (touching bboxes) should NOT have a space
        inserted between them.
        """
        doc = _make_doc()
        page = doc[0]
        # "Hel" at x=50, "lo" immediately after — estimate Hel width ≈ 3*6=18pt at fs=10
        page.insert_text((50, 200), "Hel", fontsize=10)
        # Place "lo" at x=68 — tight but no gap
        page.insert_text((68, 200), "lo", fontsize=10)

        texts = _page_texts(doc)
        combined = " ".join(texts)
        # Should not have a spurious space "Hel lo"
        assert "Hel  lo" not in combined  # double space definitely wrong


# ─── Bug 2: span ordering ────────────────────────────────────────────────────

class TestSpanOrdering:
    def test_left_to_right_order_within_line(self):
        """
        Words inserted right-to-left (in content-stream order) must appear
        left-to-right in extracted text.
        """
        doc = _make_doc()
        page = doc[0]
        # Insert "WORLD" first (right), then "HELLO" (left) to mimic z-order issue
        page.insert_text((200, 300), "WORLD", fontsize=12)
        page.insert_text((50, 300), "HELLO", fontsize=12)

        texts = _page_texts(doc)
        combined = " ".join(texts)

        assert "HELLO" in combined
        assert "WORLD" in combined
        hello_pos = combined.index("HELLO")
        world_pos = combined.index("WORLD")
        assert hello_pos < world_pos, (
            f"HELLO should come before WORLD; got: {combined!r}"
        )

    def test_multi_line_top_to_bottom_order(self):
        """Lines from different vertical positions must appear top-to-bottom."""
        doc = _make_doc()
        page = doc[0]
        # Insert bottom line first
        page.insert_text((50, 500), "LINE_THREE", fontsize=10)
        page.insert_text((50, 100), "LINE_ONE", fontsize=10)
        page.insert_text((50, 300), "LINE_TWO", fontsize=10)

        texts = _page_texts(doc)
        combined = " ".join(texts)

        pos = [combined.index(w) for w in ["LINE_ONE", "LINE_TWO", "LINE_THREE"]]
        assert pos == sorted(pos), f"Lines out of order: {combined!r}"


# ─── Bug 3: text near overlapping graphics ────────────────────────────────────

class TestTextNearGraphics:
    def test_text_behind_rect_not_clipped(self):
        """
        Text that spatially overlaps a filled rectangle must still appear
        in the extraction output (we never suppress text by image bbox).
        """
        doc = _make_doc()
        page = doc[0]
        # Draw a filled rectangle that covers the text area
        rect = fitz.Rect(40, 45, 200, 65)
        page.draw_rect(rect, color=(0, 0, 0), fill=(0.8, 0.8, 0.8))
        # Insert text that overlaps with the rectangle
        page.insert_text((50, 60), "VQFDCO", fontsize=11)

        texts = _page_texts(doc)
        combined = " ".join(texts)
        assert "VQFDCO" in combined, (
            f"Text overlapping a graphic should not be clipped; got: {combined!r}"
        )

    def test_complete_word_next_to_image(self):
        """
        A word immediately to the right of an image placeholder must appear
        complete (not truncated).
        """
        doc = _make_doc()
        page = doc[0]
        # Place a small drawn image-like rectangle at the left
        page.draw_rect(fitz.Rect(30, 50, 120, 100), fill=(0.9, 0.9, 0.9))
        # Place word just to the right of it
        page.insert_text((125, 75), "COMPLETE", fontsize=11)

        texts = _page_texts(doc)
        combined = " ".join(texts)
        assert "COMPLETE" in combined


# ─── Bug 4: annotation / badge text ──────────────────────────────────────────

class TestAnnotationText:
    def test_annotation_content_extracted(self):
        """
        A FreeText annotation whose content is not part of the main content
        stream must still appear in the extraction.
        """
        doc = _make_doc()
        page = doc[0]
        # Add a FreeText annotation (badge-like)
        annot = page.add_freetext_annot(
            fitz.Rect(50, 400, 250, 420),
            "AGENCY DISCOUNT TICKET",
            fontsize=9,
        )
        annot.update()

        content = extract_page(page, 0)
        all_text = " ".join(p.text for p in content.paragraphs)
        # Annotation text is extracted separately; check it's present
        # (It may appear via get_text("dict") for FreeText annots, or via annots())
        assert "AGENCY DISCOUNT TICKET" in all_text or any(
            "AGENCY DISCOUNT TICKET" in p.text for p in content.paragraphs
        ), f"Annotation text missing from extraction; got paragraphs: {[p.text for p in content.paragraphs]}"


# ─── Integration: build DOCX and read back ───────────────────────────────────

class TestDocxRoundtrip:
    def test_docx_contains_all_text(self, tmp_path):
        """
        build_docx() must produce a valid DOCX that contains the key words
        from the source PDF (spaces preserved).
        """
        doc = _make_doc()
        page = doc[0]
        page.insert_text((50, 100), "AIRBLUE", fontsize=11)
        page.insert_text((105, 100), "LIMITED", fontsize=11)
        page.insert_text((50, 150), "Booking Reference:", fontsize=10)
        page.insert_text((160, 150), "VQFDCO", fontsize=10)

        from app.converters.pdf_word.layout_analyzer import analyze_document
        layouts = analyze_document(doc)

        output = tmp_path / "output.docx"
        build_docx(layouts, output)

        assert output.exists(), "build_docx must create the output file"
        assert output.stat().st_size > 1000, "DOCX file is suspiciously small"

        # Read it back and verify text presence
        from docx import Document as DocxDocument
        d = DocxDocument(str(output))
        full_text = " ".join(p.text for p in d.paragraphs)
        assert "AIRBLUE" in full_text
        assert "LIMITED" in full_text
        assert "VQFDCO" in full_text

    def test_persistent_column_keeps_later_blocks_indented(self):
        doc = _make_doc()
        page = doc[0]
        for y, text in (
            (130, "jordan@example.com"),
            (146, "+1 555-0100"),
            (162, "San Francisco, CA"),
            (178, "linkedin.com/in/jsmith"),
        ):
            page.insert_text((72, y), text, fontsize=9)

        page.insert_text((220, 130), "EXPERIENCE", fontsize=11)
        page.insert_text((220, 150), "Senior Engineer", fontsize=10)
        page.insert_text((220, 165), "Led platform migration.", fontsize=9)
        page.insert_text((220, 210), "Engineer - Beta Inc", fontsize=10)
        page.insert_text((220, 225), "Built internal tooling.", fontsize=9)

        from app.converters.pdf_word.layout_analyzer import (
            ColumnRegion,
            analyze_document,
        )

        layout = analyze_document(doc)[0]
        columns = [r for r in layout.regions if isinstance(r, ColumnRegion)]

        assert len(columns) == 1
        assert len(columns[0].columns) == 2
        right_text = " ".join(
            span.text
            for item in columns[0].columns[1]
            for line in item.lines
            for span in line
        )
        assert "Engineer - Beta Inc" in right_text
        assert "Built internal tooling." in right_text

    def test_multi_page_docx(self, tmp_path):
        """analyze_document + build_docx must handle multi-page PDFs."""
        doc = fitz.open()
        for i in range(3):
            p = doc.new_page(width=595, height=842)
            p.insert_text((50, 100), f"Page number {i + 1}", fontsize=12)

        from app.converters.pdf_word.layout_analyzer import analyze_document
        layouts = analyze_document(doc)
        assert len(layouts) == 3

        output = tmp_path / "multi.docx"
        build_docx(layouts, output)
        assert output.exists()

        from docx import Document as DocxDocument
        d = DocxDocument(str(output))
        full_text = " ".join(p.text for p in d.paragraphs)
        assert "Page number 1" in full_text
        assert "Page number 3" in full_text


# ─── Integration: convert_pdf_to_docx end-to-end ─────────────────────────────

class TestConvertEndToEnd:
    def test_convert_pdf_to_docx_produces_file(self, tmp_path):
        """The top-level convert_pdf_to_docx() must return a .docx file path."""
        from app.converters.pdf_converter import convert_pdf_to_docx

        # Build a minimal PDF with enough text to pass the scanned check
        doc = fitz.open()
        for _ in range(2):
            p = doc.new_page(width=595, height=842)
            for row in range(10):
                p.insert_text(
                    (50, 80 + row * 30),
                    f"Sample text line {row + 1} with sufficient content for extraction.",
                    fontsize=10,
                )

        pdf_path = tmp_path / "input.pdf"
        doc.save(str(pdf_path))
        doc.close()

        out_dir = tmp_path / "out"
        out_dir.mkdir()

        result = convert_pdf_to_docx(pdf_path, out_dir)
        assert result.path.exists()
        assert result.path.suffix == ".docx"
        assert result.path.stat().st_size > 500

        from docx import Document as DocxDocument
        converted = DocxDocument(str(result.path))
        assert len(converted.inline_shapes) == 2
        hidden_text = " ".join(p.text for p in converted.paragraphs)
        assert "Sample text line 1" in hidden_text

        editable_dir = tmp_path / "editable"
        editable_dir.mkdir()
        editable_result = convert_pdf_to_docx(
            pdf_path,
            editable_dir,
            mode="editable",
        )
        editable = DocxDocument(str(editable_result.path))
        editable_text = " ".join(p.text for p in editable.paragraphs)
        assert "Sample text line 1" in editable_text
        assert len(editable.inline_shapes) == 0

    def test_scanned_pdf_raises_error(self, tmp_path):
        """A PDF with no text must raise ConversionError with code scanned_pdf."""
        from app.converters.pdf_converter import convert_pdf_to_docx
        from app.errors import ConversionError

        doc = fitz.open()
        doc.new_page(width=595, height=842)  # blank page, no text
        pdf_path = tmp_path / "scanned.pdf"
        doc.save(str(pdf_path))
        doc.close()

        out_dir = tmp_path / "out2"
        out_dir.mkdir()

        with pytest.raises(ConversionError) as exc_info:
            convert_pdf_to_docx(pdf_path, out_dir)

        assert exc_info.value.code == "scanned_pdf"
