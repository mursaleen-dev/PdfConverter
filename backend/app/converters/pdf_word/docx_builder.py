"""
DOCX writer that consumes PageLayout produced by layout_analyzer.py.

Handles each region type:
  TableRegion   → python-docx Table with borders and optional cell shading
  BoxRegion     → single-cell bordered/shaded table (badge, notice box)
  ColumnRegion  → borderless 2-column table
  ParagraphRegion → styled paragraph with per-span runs
  ImageRegion   → inline picture scaled to fit page
"""
from __future__ import annotations

import io
import logging
from pathlib import Path
from typing import Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches, Emu

from .layout_analyzer import (
    PageLayout,
    TableRegion, BoxRegion, ParagraphRegion, ColumnRegion, ImageRegion,
    Span,
)

logger = logging.getLogger(__name__)

_MAX_IMAGE_WIDTH_IN = 6.0
_KEEP_LAYOUT_DPI = 144
_HEADING_SIZE_RATIO = 1.3   # dom_size / body_size ratio → heading

_PT_TO_EMU = 12700
_PAGE_MARGIN_IN = 0.35  # narrow margin so reflowed text wraps close to the
                        # source PDF's line breaks instead of at python-docx's
                        # much wider 1.25in default, which was inflating
                        # content height and pushing single-page PDFs to
                        # spill onto a second page in the output.


# ─── color helpers ────────────────────────────────────────────────────────────

def _rgb_from_int(packed: int) -> Optional[RGBColor]:
    """Convert packed 0xRRGGBB PDF color to RGBColor, preserving more colors."""
    if packed == 0:
        return None
    r = (packed >> 16) & 0xFF
    g = (packed >> 8)  & 0xFF
    b =  packed        & 0xFF

    # Only drop pure black and pure white
    if (r, g, b) == (0, 0, 0):
        return None  # pure black → use default
    if (r, g, b) == (255, 255, 255):
        return None  # pure white → use default

    # Keep all other colors including near-black (better preserves blue/gray text)
    return RGBColor(r, g, b)


# ─── XML helpers for tables ───────────────────────────────────────────────────

def _set_cell_borders(cell, color_hex: str = "000000", size: int = 4) -> None:
    """Set all four borders of a table cell to a thin single line."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    str(size))
        el.set(qn("w:color"), color_hex)
        el.set(qn("w:space"), "0")
        tcBorders.append(el)
    tcPr.append(tcBorders)


def _set_cell_no_borders(cell) -> None:
    """Remove all borders from a table cell (for borderless column layouts)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "none")
        tcBorders.append(el)
    tcPr.append(tcBorders)


def _set_cell_shading(cell, fill_hex: str) -> None:
    """Set cell background color via XML shading element."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill_hex.upper())
    tcPr.append(shd)


def _set_table_no_spacing(table) -> None:
    """Optimize table spacing and appearance."""
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)

    # Zero out cell spacing
    spacing = OxmlElement("w:tblCellSpacing")
    spacing.set(qn("w:w"),    "0")
    spacing.set(qn("w:type"), "dxa")
    tblPr.append(spacing)

    # Shrink Word's default cell padding (~110 twips top/bottom) — the source
    # PDF's rows are packed much tighter, and the default padding was one of
    # the contributors to reconstructed tables running taller than the
    # original and pushing trailing content onto an extra page.
    cell_mar = OxmlElement("w:tblCellMar")
    for edge, twips in (("top", "10"), ("bottom", "10"), ("left", "60"), ("right", "60")):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:w"), twips)
        el.set(qn("w:type"), "dxa")
        cell_mar.append(el)
    tblPr.append(cell_mar)


def _set_table_geometry(table, widths_pt: list[float],
                        left_indent_pt: float = 0.0) -> None:
    """Use source-PDF geometry instead of Word's equal-width autofit."""
    if not widths_pt:
        return
    table.autofit = False
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)

    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblPr.append(layout)

    indent = OxmlElement("w:tblInd")
    indent.set(qn("w:w"), str(max(0, round(left_indent_pt * 20))))
    indent.set(qn("w:type"), "dxa")
    tblPr.append(indent)

    grid = tbl.tblGrid
    for grid_col, width_pt in zip(grid.gridCol_lst, widths_pt):
        grid_col.set(qn("w:w"), str(max(1, round(width_pt * 20))))

    for row in table.rows:
        for cell, width_pt in zip(row.cells, widths_pt):
            cell.width = Pt(max(width_pt, 1.0))


# ─── span grouping helper ─────────────────────────────────────────────────────

def _group_into_lines(spans: list[Span]) -> list[list[Span]]:
    """Cluster spans by y-center proximity into visual lines, sorted left-to-right."""
    lines: list[list[Span]] = []
    for span in sorted(spans, key=lambda s: (s.y0, s.x0)):
        placed = False
        for line in reversed(lines):
            rep_y = sum(s.y_center for s in line) / len(line)
            if abs(span.y_center - rep_y) <= 3.0:
                line.append(span)
                placed = True
                break
        if not placed:
            lines.append([span])
    for line in lines:
        line.sort(key=lambda s: s.x0)
    return lines


# ─── run builder ─────────────────────────────────────────────────────────────

def _spans_to_runs(para, spans: list[Span], dom_size: float,
                   is_heading: bool = False) -> None:
    """Append styled DOCX runs for spans, synthesising spaces where needed.

    Icon spans (rasterized icon-font glyphs) become their own inline
    picture run instead of being merged into the text-grouping below.
    """
    if not spans:
        return

    # Build token list with synthesised spaces for visual word gaps
    tokens: list[Span] = []
    for si, span in enumerate(spans):
        tokens.append(span)
        if si < len(spans) - 1:
            nxt = spans[si + 1]
            if span.is_icon or nxt.is_icon:
                continue  # icons carry their own visual padding; no synthesized space
            gap = nxt.x0 - span.x1
            avg_cw = max(1.0, span.size * 0.5)
            if (gap >= 0.15 * avg_cw
                    and not span.text.endswith(" ")
                    and not nxt.text.startswith(" ")):
                tokens.append(Span(
                    text=" ", bold=span.bold, italic=span.italic,
                    size=span.size, color=span.color,
                    x0=span.x1, y0=span.y0, x1=nxt.x0, y1=span.y1,
                ))

    # Merge consecutive same-style TEXT tokens into one run; icon tokens
    # always become their own inline picture run and break the merge chain.
    groups: list[tuple] = []
    for tok in tokens:
        if tok.is_icon:
            groups.append(("icon", tok.icon_data, tok.x1 - tok.x0, tok.y1 - tok.y0))
            continue
        key = (tok.bold, tok.italic, tok.color, tok.size)
        if groups and groups[-1][0] == "text" and groups[-1][2:] == key:
            groups[-1] = ("text", groups[-1][1] + tok.text, *key)
        else:
            groups.append(("text", tok.text, tok.bold, tok.italic, tok.color, tok.size))

    for group in groups:
        if group[0] == "icon":
            _, icon_data, w_pt, h_pt = group
            if not icon_data or h_pt <= 0:
                continue
            try:
                run = para.add_run()
                run.add_picture(io.BytesIO(icon_data), height=Pt(max(h_pt, 1.0)))
            except Exception:
                logger.debug("failed to insert inline icon picture")
            continue

        _, text, bold, italic, color, size = group
        run = para.add_run(text)
        if bold:
            run.bold = True
        elif not is_heading:
            run.bold = False
        if italic:
            run.italic = True
        elif not is_heading:
            run.italic = False
        run.font.size = Pt(size)
        col = _rgb_from_int(color)
        if col:
            run.font.color.rgb = col


# ─── region writers ───────────────────────────────────────────────────────────

def _add_paragraph_region(doc: Document, region: ParagraphRegion,
                           body_size: float) -> None:
    if not region.lines:
        return

    dom_size = region.dominant_size()
    is_heading = dom_size >= body_size * _HEADING_SIZE_RATIO

    # Add each line as a separate paragraph to preserve line breaks and spacing
    for li, line in enumerate(region.lines):
        if not line:
            continue

        if is_heading and li == 0:
            ratio = dom_size / body_size
            level = 1 if ratio >= 2.0 else (2 if ratio >= 1.6 else 3)
            wp = doc.add_heading("", level=level)
        else:
            wp = doc.add_paragraph("", style="Normal")

        _spans_to_runs(wp, line, dom_size, is_heading=(is_heading and li == 0))

        # PDF lines already carry their own y geometry. Word's default
        # paragraph after-spacing compounds across every extracted line and
        # pushes lower page regions down, so keep adjacent source lines tight.
        wp.paragraph_format.line_spacing = 1.0
        wp.paragraph_format.space_before = Pt(0)
        wp.paragraph_format.space_after = Pt(0)


def _tighten_paragraph(p) -> None:
    """Zero out a paragraph's spacing/line-height. Table/box/column cell
    paragraphs otherwise inherit Normal style's default spacing (commonly
    ~8pt space-after with ~1.08x line height), which — multiplied across
    every row of every table — was a large, easy-to-miss contributor to
    reconstructed documents running taller than the source PDF."""
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0


def _add_table_region(doc: Document, region: TableRegion,
                      body_size: float) -> None:
    n_rows = len(region.rows)
    n_cols = max((len(row) for row in region.rows), default=0)
    if n_rows == 0 or n_cols == 0:
        return

    wt = doc.add_table(rows=n_rows, cols=n_cols)
    wt.style = "Table Grid"
    _set_table_no_spacing(wt)
    if len(region.col_xs) >= n_cols + 1:
        widths = [
            max(region.col_xs[i + 1] - region.col_xs[i], 1.0)
            for i in range(n_cols)
        ]
        _set_table_geometry(
            wt,
            widths,
            left_indent_pt=region.col_xs[0] - (_PAGE_MARGIN_IN * 72.0),
        )

    for ri, row in enumerate(region.rows):
        for ci in range(n_cols):
            cell = wt.cell(ri, ci)
            _set_cell_borders(cell, color_hex="333333", size=4)

            if ci >= len(row):
                continue
            tc_data = row[ci]

            if tc_data.bg_hex:
                _set_cell_shading(cell, tc_data.bg_hex)

            if not tc_data.spans:
                continue

            lines = _group_into_lines(tc_data.spans)
            wp = cell.paragraphs[0]
            for li, line in enumerate(lines):
                if li > 0:
                    wp = cell.add_paragraph()
                _tighten_paragraph(wp)
                _spans_to_runs(wp, line, body_size)


def _add_box_region(doc: Document, region: BoxRegion, body_size: float) -> None:
    """Render as a single-cell table with optional background and border."""
    wt = doc.add_table(rows=1, cols=1)
    _set_table_no_spacing(wt)

    cell = wt.cell(0, 0)
    border_hex = region.border_hex or "333333"
    _set_cell_borders(cell, color_hex=border_hex, size=6)

    if region.bg_hex:
        _set_cell_shading(cell, region.bg_hex)

    if not region.spans:
        return

    lines = _group_into_lines(region.spans)
    wp = cell.paragraphs[0]
    for li, line in enumerate(lines):
        if li > 0:
            wp = cell.add_paragraph()
        _tighten_paragraph(wp)
        _spans_to_runs(wp, line, body_size)


def _add_column_region(doc: Document, region: ColumnRegion,
                       body_size: float) -> None:
    n_cols = len(region.columns)
    if n_cols == 0:
        return

    wt = doc.add_table(rows=1, cols=n_cols)
    _set_table_no_spacing(wt)

    boxes: list[tuple[float, float]] = []
    for col_items in region.columns:
        x0s: list[float] = []
        x1s: list[float] = []
        for item in col_items:
            if isinstance(item, ImageRegion):
                x0s.append(item.x0)
                x1s.append(item.x1)
            else:
                spans = [span for line in item.lines for span in line]
                x0s.extend(span.x0 for span in spans)
                x1s.extend(span.x1 for span in spans)
        boxes.append((min(x0s), max(x1s)) if x0s else (0.0, 1.0))

    if boxes:
        boundaries = [boxes[0][0]]
        for left, right in zip(boxes, boxes[1:]):
            boundaries.append((left[1] + right[0]) / 2)
        boundaries.append(boxes[-1][1])
        widths = [
            max(boundaries[i + 1] - boundaries[i], 1.0)
            for i in range(n_cols)
        ]
        _set_table_geometry(
            wt,
            widths,
            left_indent_pt=boundaries[0] - (_PAGE_MARGIN_IN * 72.0),
        )

    for ci, col_items in enumerate(region.columns):
        cell = wt.cell(0, ci)
        _set_cell_no_borders(cell)

        first_para = True
        for item in col_items:
            if isinstance(item, ImageRegion):
                wp = cell.paragraphs[0] if first_para else cell.add_paragraph()
                first_para = False
                _tighten_paragraph(wp)
                _add_image_to_paragraph(wp, item, n_cols)
                continue

            for li, line in enumerate(item.lines):
                if first_para and li == 0:
                    wp = cell.paragraphs[0]
                else:
                    wp = cell.add_paragraph()
                first_para = False
                _tighten_paragraph(wp)
                _spans_to_runs(wp, line, body_size)


_THIN_DECOR_HEIGHT_PT = 6.0  # source-PDF height below which an image region is
                              # treated as a hairline divider/decoration rather
                              # than a real picture, and given ~zero paragraph
                              # overhead instead of a full text line's worth.


def _add_image_region(doc: Document, region: ImageRegion) -> None:
    if not region.data or len(region.data) == 0:
        return
    w_pt = region.x1 - region.x0
    h_pt = region.y1 - region.y0
    if w_pt <= 0:
        return
    try:
        is_thin_decor = h_pt <= _THIN_DECOR_HEIGHT_PT

        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0 if is_thin_decor else 6)

        stream = io.BytesIO(region.data)
        desired_w_in = min(w_pt / 72.0, _MAX_IMAGE_WIDTH_IN)

        run = p.add_run()
        run.add_picture(stream, width=Inches(desired_w_in))

        if is_thin_decor:
            # Pin the paragraph's line height to the (tiny) scaled image
            # height instead of inheriting Normal style's full text-line
            # height, which would otherwise reserve a whole line of vertical
            # space for a 1-6pt-tall divider/decoration.
            scaled_h_pt = max(h_pt * (desired_w_in * 72.0 / max(w_pt, 1.0)), 1.0)
            pPr = p._p.get_or_add_pPr()
            spacing = pPr.find(qn("w:spacing"))
            if spacing is None:
                spacing = OxmlElement("w:spacing")
                pPr.append(spacing)
            spacing.set(qn("w:line"), str(int(scaled_h_pt * 20)))  # twentieths of a pt
            spacing.set(qn("w:lineRule"), "exact")

        logger.debug("Added image xref=%d, size=%d bytes, width=%.2fin",
                     region.xref, len(region.data), desired_w_in)
    except Exception as exc:
        logger.warning("Failed to add image xref=%d: %s", region.xref, exc)


def _add_image_to_paragraph(p, region: ImageRegion, n_cols: int) -> None:
    """Insert an image inline into an existing (column-cell) paragraph —
    used when an image sits beside text as a column (e.g. a barcode next to
    its label), rather than as its own standalone block."""
    if not region.data:
        return
    w_pt = region.x1 - region.x0
    if w_pt <= 0:
        return
    try:
        # Cap width to roughly this column's share of the page so a wide
        # source image (e.g. a barcode) doesn't blow out the cell width.
        max_col_w_in = _MAX_IMAGE_WIDTH_IN / max(1, n_cols)
        desired_w_in = min(w_pt / 72.0, max_col_w_in)
        run = p.add_run()
        run.add_picture(io.BytesIO(region.data), width=Inches(desired_w_in))
    except Exception as exc:
        logger.warning("Failed to add column image xref=%d: %s", region.xref, exc)


# ─── page break ───────────────────────────────────────────────────────────────

def _insert_page_break(doc: Document) -> None:
    p  = OxmlElement("w:p")
    r  = OxmlElement("w:r")
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    r.append(br)
    p.append(r)
    doc.element.body.append(p)


# ─── body-size estimation ─────────────────────────────────────────────────────

def _body_font_size(layouts: list[PageLayout]) -> float:
    sizes: list[float] = []
    for layout in layouts:
        for region in layout.regions:
            if isinstance(region, ParagraphRegion):
                for line in region.lines:
                    for span in line:
                        if span.text.strip():
                            sizes.append(span.size)
    if not sizes:
        return 10.0
    sizes.sort()
    return sizes[len(sizes) // 2]


# ─── main entry point ─────────────────────────────────────────────────────────

def build_docx(layouts: list[PageLayout], output_path: Path) -> None:
    """
    Write a DOCX file from a list of PageLayout objects produced by
    layout_analyzer.analyze_document().
    """
    doc = Document()

    # Remove default content (paragraphs and tables) but keep sectPr,
    # which python-docx needs for page-layout queries like _block_width.
    body = doc.element.body
    for el in list(body):
        if el.tag in (qn("w:p"), qn("w:tbl")):
            body.remove(el)

    # Match the generated page size to the source PDF instead of using
    # python-docx's hardcoded US-Letter default. A mismatched (typically
    # narrower) page width changes where text wraps versus the original,
    # which inflates content height enough to push single-page PDFs onto a
    # second page. Margins are narrowed too, for the same reason.
    if layouts:
        section = doc.sections[0]
        section.page_width = Emu(round(layouts[0].width * _PT_TO_EMU))
        section.page_height = Emu(round(layouts[0].height * _PT_TO_EMU))
        section.left_margin = Inches(_PAGE_MARGIN_IN)
        section.right_margin = Inches(_PAGE_MARGIN_IN)
        section.top_margin = Inches(_PAGE_MARGIN_IN)
        section.bottom_margin = Inches(_PAGE_MARGIN_IN)

    body_size = _body_font_size(layouts)

    for page_idx, layout in enumerate(layouts):
        if page_idx > 0:
            _insert_page_break(doc)

        prev_region_y1 = None
        prev_was_thin_decor = False
        for region in layout.regions:
            is_thin_decor = (
                isinstance(region, ImageRegion)
                and (region.y1 - region.y0) <= _THIN_DECOR_HEIGHT_PT
            )

            # Add spacing if there's a gap from previous region. Skip this
            # around thin decorative images (divider lines etc.) — they're
            # meant to sit flush against surrounding content, and with N of
            # them on a page the extra gap paragraphs added up to enough
            # height to push content onto an extra page.
            if prev_region_y1 is not None and not is_thin_decor and not prev_was_thin_decor:
                region_y0 = getattr(region, 'y0', None)
                if region_y0 is not None and region_y0 - prev_region_y1 > 12:
                    # Add spacing paragraph for visual separation
                    spacing_para = doc.add_paragraph()
                    spacing_para.paragraph_format.space_before = Pt(3)
                    spacing_para.paragraph_format.space_after = Pt(3)

            prev_was_thin_decor = is_thin_decor

            if isinstance(region, TableRegion):
                _add_table_region(doc, region, body_size)
                prev_region_y1 = region.y0 + sum(len(row) for row in region.rows) * 12
            elif isinstance(region, BoxRegion):
                _add_box_region(doc, region, body_size)
                prev_region_y1 = region.y1
            elif isinstance(region, ColumnRegion):
                _add_column_region(doc, region, body_size)
                if region.columns:
                    max_lines = max(len(col) for col in region.columns)
                    prev_region_y1 = region.y0 + max_lines * 12
                else:
                    prev_region_y1 = region.y0
            elif isinstance(region, ImageRegion):
                _add_image_region(doc, region)
                prev_region_y1 = region.y1
            elif isinstance(region, ParagraphRegion):
                _add_paragraph_region(doc, region, body_size)
                if region.lines:
                    prev_region_y1 = region.y0 + len(region.lines) * body_size * 0.3
                else:
                    prev_region_y1 = region.y0
            else:
                logger.debug("Unknown region type: %s", type(region))

    # Word requires a paragraph between the last table and the section
    # properties (it silently inserts one on open if missing). Our documents
    # often end with a table/box region, and since content is now packed
    # close to the page edge, Word's auto-inserted paragraph was enough to
    # spill onto a blank trailing page — even though LibreOffice tolerates
    # the missing paragraph and renders one page. Add an explicit
    # zero-height paragraph so nothing gets added implicitly.
    trailing = doc.add_paragraph()
    trailing.paragraph_format.space_before = Pt(0)
    trailing.paragraph_format.space_after = Pt(0)
    tpPr = trailing._p.get_or_add_pPr()
    tspacing = OxmlElement("w:spacing")
    tspacing.set(qn("w:line"), "20")  # 1pt, twentieths-of-a-point units
    tspacing.set(qn("w:lineRule"), "exact")
    tpPr.append(tspacing)
    run = trailing.add_run("")
    run.font.size = Pt(1)

    doc.save(str(output_path))
    logger.info("DOCX written to %s (%d pages)", output_path, len(layouts))


def build_keep_layout_docx(pdf_doc, output_path: Path) -> None:
    """Write a visually exact DOCX using one rendered image per PDF page.

    A hidden text run keeps source text available to indexing, search, and
    downstream extraction without affecting the rendered page geometry.
    """
    doc = Document()
    body = doc.element.body
    for el in list(body):
        if el.tag in (qn("w:p"), qn("w:tbl")):
            body.remove(el)

    if pdf_doc.page_count == 0:
        doc.save(str(output_path))
        return

    first_page = pdf_doc[0]
    section = doc.sections[0]
    section.page_width = Emu(round(first_page.rect.width * _PT_TO_EMU))
    section.page_height = Emu(round(first_page.rect.height * _PT_TO_EMU))
    section.left_margin = Emu(0)
    section.right_margin = Emu(0)
    section.top_margin = Emu(0)
    section.bottom_margin = Emu(0)
    section.header_distance = Emu(0)
    section.footer_distance = Emu(0)

    for page_index, page in enumerate(pdf_doc):
        paragraph = doc.add_paragraph()
        if page_index:
            paragraph.paragraph_format.page_break_before = True
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.left_indent = Pt(0)
        paragraph.paragraph_format.right_indent = Pt(0)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        pix = page.get_pixmap(dpi=_KEEP_LAYOUT_DPI, alpha=False)
        image = io.BytesIO(pix.tobytes("png"))

        # Inline pictures need a little room for Word's mandatory paragraph
        # mark. Scale proportionally instead of forcing an exact line height:
        # exact line height clips the top of tall images in LibreOffice/Word.
        image_height = max(page.rect.height - 8.0, 1.0)
        image_width = page.rect.width * (image_height / page.rect.height)
        run = paragraph.add_run()
        run.add_picture(
            image,
            width=Pt(image_width),
            height=Pt(image_height),
        )

        hidden = paragraph.add_run(page.get_text("text"))
        hidden.font.hidden = True
        hidden.font.size = Pt(1)

    doc.save(str(output_path))
    logger.info(
        "Keep-layout DOCX written to %s (%d pages)",
        output_path,
        pdf_doc.page_count,
    )
